import json
import re
import asyncio
from typing import List, Dict, Any
from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, JSONResponse
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
import io
import os

# Import Models
from models import format_time

# Import New Architecture Components
from llm_gateway import LLMGateway
from agents.user_agent import UserContextAgent
from agents.discovery_agent import JobDiscoveryAgent
from agents.matching_agent import MatchingAgent
from cv_service import analyze_cv

try:
    from ml_model.salary_predictor import load_model, predict_missing_salary, train_and_save, MODEL_PATH_DEFAULT
    from ml_model.prepare_training_data import prepare_training_data
    salary_model = load_model()
    if salary_model:
        print(f"[OK] Salary prediction model loaded ({MODEL_PATH_DEFAULT})")
    else:
        print(f"[INFO] Salary prediction model not found (run training to create {MODEL_PATH_DEFAULT})")
except Exception as e:
    print("[WARN] Salary predictor unavailable:", e)
    salary_model = None

import pandas as pd

app = FastAPI()

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/")
async def start():
    with open("static/index.html") as f:
        return HTMLResponse(content=f.read())

# New helper: parse uploaded CV file (pdf, docx, txt)
async def _parse_cv_file(file: UploadFile) -> Dict[str, Any]:
    """
    Returns dict: {"status": "ok", "text": "...", "file_type": "pdf|docx|txt"}
    or {"status": "error", "error": "msg"} on failure.
    """
    filename = file.filename.lower()
    content_type = file.content_type
    text = ""

    try:
        if filename.endswith(".pdf"):
            if not pdfplumber:
                return {"status": "error", "error": "PDF support not installed (pip install pdfplumber)"}
            try:
                # file.file is a SpooledTemporaryFile-like object
                with pdfplumber.open(file.file) as pdf:
                    for page in pdf.pages:
                        extract = page.extract_text()
                        if extract:
                            text += extract + "\n"
                return {"status": "ok", "text": text.strip(), "file_type": "pdf"}
            except Exception as e:
                return {"status": "error", "error": f"Invalid PDF file: {e}"}

        if filename.endswith(".docx"):
            try:
                import docx
            except Exception:
                return {"status": "error", "error": "DOCX support not installed (pip install python-docx)"}
            try:
                content_bytes = await file.read()
                doc = docx.Document(io.BytesIO(content_bytes))
                text = "\n".join([p.text for p in doc.paragraphs if p.text])
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text:
                                text += "\n" + cell.text
                return {"status": "ok", "text": text.strip(), "file_type": "docx"}
            except Exception as e:
                return {"status": "error", "error": f"Invalid DOCX file: {e}"}

        if filename.endswith(".txt") or content_type == "text/plain":
            try:
                content = await file.read()
                text = content.decode("utf-8")
                return {"status": "ok", "text": text.strip(), "file_type": "txt"}
            except Exception as e:
                return {"status": "error", "error": f"Invalid TXT file: {e}"}

        return {"status": "error", "error": "Only .pdf, .docx, and .txt files are supported currently"}
    except Exception as e:
        return {"status": "error", "error": str(e)}

@app.post("/upload_cv")
async def upload_cv(file: UploadFile = File(...)):
    print(f"[INFO] CV upload received: {file.filename}")

    try:
        parsed = await _parse_cv_file(file)
        if parsed.get("status") != "ok":
            return JSONResponse({"status": "error", "error": parsed.get("error")}, status_code=400)

        text = parsed.get("text", "")
        file_type = parsed.get("file_type", "unknown")

        if not text:
            return JSONResponse({"status": "error", "error": "Could not read text from this file. If this is a scanned PDF (image), please convert it to text first."}, status_code=400)

        # return both text and the detected file type so frontend can display / store metadata
        return {"status": "success", "file_type": file_type, "text": text}

    except Exception as e:
        print(f"[ERROR] Upload error: {e}")
        return JSONResponse({"status": "error", "error": str(e)}, status_code=500)

# --- Initialization ---
print("[INFO] INITIALIZING TwinWork AI Agents...")

# 0. Infrastructure
llm_gateway = LLMGateway()

# 1. Agents
discovery_agent = JobDiscoveryAgent()
matching_agent = MatchingAgent()

# Note: UserContextAgent is stateful per session, so we init it in the WebSocket handler

print("[OK] Agents Ready")

class ChatSession:
    def __init__(self, websocket: WebSocket):
        self.websocket = websocket
        self.user_id = str(id(websocket))
        
        # Initialize User Agent for this session
        self.user_agent = UserContextAgent(llm_gateway)
        
        self.found_jobs = []

    async def send_message(self, message: str, type: str = "text", options: List[str] = None):
        await self.websocket.send_json({
            "type": type,
            "message": message,
            "options": options
        })

    async def handle_message(self, raw_message: str):
        try:
            # Parse JSON if possible
            data = json.loads(raw_message)
            text = ""
            if isinstance(data, dict):
                text = data.get("text", data.get("message", ""))
                # Handle types directly if needed (e.g. feedback)
                if data.get("type") == "feedback":
                    # TODO: Implement feedback handling via Agent
                    return
                if data.get("type") == "cv_upload":
                    cv_text = data.get("content", "")
                    await self.handle_cv_upload(cv_text)
                    return
            else:
                text = str(data)
                
            await self.process_input(text)
            
        except json.JSONDecodeError:
            await self.process_input(raw_message)

    async def handle_cv_upload(self, cv_text: str):
        """Process uploaded CV text and update agent state"""
        await self.send_message("[INFO] Analyzing your CV... this might take a moment.")
        
        try:
            # 1. Analyze with CV Service
            analysis = await analyze_cv(cv_text)
            cv_data = analysis.get('cv_data', {})
            skills_raw = cv_data.get('skills', [])
            # Normalize/tokenize skills server-side to ensure consistent matching
            skills = _tokenize_skills(skills_raw, text_fallback=cv_text)
            name = cv_data.get('name')
            experience = cv_data.get('experience', [])

            # 2. Update User Agent Profile
            # Create a synthetic "extraction" result to feed into the agent
            extracted_info = {
                "skills": skills,
                "experience_years": cv_data.get('total_experience_years'),
                "location": cv_data.get('location'),
                "cv_summary": cv_data.get('summary')
            }
            if name:
                extracted_info["name"] = name
                
            # Safely update profile
            for k, v in extracted_info.items():
                if v:
                    # ensure skills stored as normalized list
                    if k == "skills" and isinstance(v, list):
                        self.user_agent.user_profile[k] = v
                    else:
                        self.user_agent.user_profile[k] = v
                    
            # 3. Formulate response
            response = f"Thanks! I've analyzed your CV."
            if name:
                response += f" Nice to meet you, {name}!"
            
            if skills:
                top_skills = ", ".join(skills[:5])
                response += f"\n\nI see you have skills in: **{top_skills}**"
                if len(skills) > 5:
                    response += f" and {len(skills)-5} others."
            
            response += "\n\nIs this correct? Would you like me to start searching for jobs based on this?"
            
            # Send results to frontend (so UI can update state if needed)
            await self.websocket.send_json({
                "type": "cv_analysis",
                "skills": skills,
                "name": name
            })
            
            await self.send_message(response)
            
        except Exception as e:
            print(f"[ERROR] Error handling CV: {e}")
            await self.send_message("I had some trouble analyzing your CV, but I received the text. Let's continue chatting!")

    async def process_input(self, text: str):
        # 1. Delegate to User Agent
        try:
            response_text, ready_to_search = await self.user_agent.process_message(text)
        except Exception as e:
            print(f"[WARN] UserAgent processing error: {e}")
            await self.send_message("Sorry, I'm having trouble generating a response right now. Try again later.")
            return
        
        # 2. Send Agent Response
        if response_text:
            await self.send_message(response_text)
            
        # 3. Check for Search Trigger
        if ready_to_search:
            await self.perform_search()

    async def perform_search(self):
        profile = self.user_agent.user_profile
        name = profile.get("name", "there")
        
        await self.send_message(f"Searching for jobs matching your profile, {name}...")
        
        # Build query and location candidates
        sub_queries = _build_query_variants(profile)
        location_variants = _build_location_candidates(profile.get("location", ""), remote_ok=bool(profile.get("remote_ok", False)))

        all_jobs = []
        seen_ids = set()
        max_calls = 60
        calls = 0

        # Strategy: try remote-tagged searches first, then per-location queries, then global
        for q in sub_queries[:10]:
            if calls >= max_calls:
                break
            # Remote-first attempt (if allowed)
            if profile.get("remote_ok", False):
                try:
                    calls += 1
                    print(f"[SEARCH] remote attempt: '{q}'")
                    sub_results = await discovery_agent.search(query=q, location="Remote", remote_only=True)
                except Exception as e:
                    print(f"[WARN] remote search failed for '{q}': {e}")
                    sub_results = []
                for job in sub_results:
                    if getattr(job, "job_id", None) and job.job_id not in seen_ids:
                        job.search_query = q
                        job.search_location = "Remote"
                        all_jobs.append(job); seen_ids.add(job.job_id)
            # Per-location attempts
            for loc in location_variants:
                if calls >= max_calls:
                    break
                # skip redundant remote token here
                if loc == "Remote":
                    continue
                try:
                    calls += 1
                    loc_label = loc or None
                    print(f"[SEARCH] '{q}' @ '{loc_label or 'GLOBAL'}' (call {calls}/{max_calls})")
                    sub_results = await discovery_agent.search(query=q, location=loc_label, remote_only=False)
                except Exception as e:
                    print(f"[WARN] search failed ({q!r},{loc!r}): {e}")
                    sub_results = []
                for job in sub_results:
                    if getattr(job, "job_id", None) and job.job_id not in seen_ids:
                        job.search_query = q
                        job.search_location = loc or "GLOBAL"
                        all_jobs.append(job); seen_ids.add(job.job_id)
                # small throttle
                await asyncio.sleep(0.03)

        # If still empty, try skill-only broad searches without location
        if not all_jobs and profile.get("skills"):
            for skill in profile.get("skills")[:6]:
                if calls >= max_calls:
                    break
                try:
                    calls += 1
                    print(f"[SEARCH] broad skill search: '{skill}'")
                    sub_results = await discovery_agent.search(query=skill, location=None, remote_only=profile.get("remote_ok", False))
                except Exception as e:
                    print(f"[WARN] broad skill search failed for '{skill}': {e}")
                    sub_results = []
                for job in sub_results:
                    if getattr(job, "job_id", None) and job.job_id not in seen_ids:
                        job.search_query = skill
                        job.search_location = "GLOBAL"
                        all_jobs.append(job); seen_ids.add(job.job_id)

        # final fallback: single generic search
        if not all_jobs:
            try:
                print("[SEARCH] final fallback: 'Software Engineer' global")
                sub_results = await discovery_agent.search(query="Software Engineer", location=None, remote_only=profile.get("remote_ok", False))
            except Exception as e:
                print(f"[WARN] fallback search failed: {e}")
                sub_results = []
            for job in sub_results:
                if getattr(job, "job_id", None) and job.job_id not in seen_ids:
                    job.search_query = "Software Engineer"
                    job.search_location = "GLOBAL"
                    all_jobs.append(job); seen_ids.add(job.job_id)

        # --- Ensure jobs variable is set for downstream code ---
        jobs = all_jobs
        
        # If still no jobs, perform a simpler location-based retry (helps with discovery providers that prefer simple queries)
        if not jobs:
            print("[SEARCH] advanced strategy returned no results — retrying simpler location-based searches")
            loc = profile.get("location", "Yerevan")
            remote_ok = bool(profile.get("remote_ok", False))
            seen_retry = set(seen_ids)
            try_queries = sub_queries[:6] if sub_queries else ["Software Engineer"]
            for q in try_queries:
                # try user's location
                try:
                    print(f"[SEARCH-RETRY] '{q}' @ '{loc}' (remote_only={remote_ok})")
                    sub_results = await discovery_agent.search(query=q, location=loc, remote_only=remote_ok)
                except Exception as e:
                    print(f"[WARN] retry search failed ({q!r},{loc!r}): {e}")
                    sub_results = []
                for job in sub_results:
                    jid = getattr(job, "job_id", None)
                    if jid and jid not in seen_retry:
                        job.search_query = q
                        job.search_location = loc
                        jobs.append(job); seen_retry.add(jid)
                # try remote explicitly if allowed
                if remote_ok:
                    try:
                        print(f"[SEARCH-RETRY] '{q}' @ 'Remote' (remote_only=True)")
                        sub_results = await discovery_agent.search(query=q, location="Remote", remote_only=True)
                    except Exception as e:
                        print(f"[WARN] remote retry failed for '{q}': {e}")
                        sub_results = []
                    for job in sub_results:
                        jid = getattr(job, "job_id", None)
                        if jid and jid not in seen_retry:
                            job.search_query = q
                            job.search_location = "Remote"
                            jobs.append(job); seen_retry.add(jid)
                if jobs:
                    break

        if not jobs:
            await self.send_message("I couldn't find any jobs right now. Try broadening your location or skills.")
            return

        self.found_jobs.extend([j for j in jobs if j.job_id not in {existing.job_id for existing in self.found_jobs}])
        await self.send_message(f"Found {len(jobs)} jobs! Analyzing best matches...")

        # 2. Matching Agent
        matches = matching_agent.analyze_matches(profile, jobs)
        
        # 3. Present Results
        single_jobs_data = []
        for job in matches.get("raw_top_jobs", []):
            # Preserve original salary (if present), compute prediction only when missing
            orig = getattr(job, "hourly_rate", None)
            orig_val = float(orig) if orig not in (None, 0, "", False) else None
            predicted_rate = None
            was_predicted = False
            if orig_val is None and salary_model:
                job_features = {
                    "job_title": getattr(job, "title", ""),
                    "required_skills": getattr(job, "skills", []),
                    "experience_level": getattr(job, "experience_level", ""),
                    "job_location": getattr(job, "location", ""),
                    "job_type": getattr(job, "job_type", ""),
                    "company_industry": getattr(job, "company_industry", "")
                }
                predicted_rate = predict_missing_salary(job_features, model=salary_model)
                was_predicted = bool(predicted_rate)
                try:
                    if was_predicted:
                        job.hourly_rate_was_predicted = True
                except Exception:
                    pass

            display_rate = orig_val if orig_val is not None else (float(predicted_rate) if predicted_rate else 0)
            salary_source = "predicted" if was_predicted else ("reported" if orig_val is not None else "unknown")

            single_jobs_data.append({
                "title": job.title,
                "company": job.company,
                "location": job.location,
                "original_hourly_rate": orig_val,
                "hourly_rate": display_rate,
                "predicted_hourly_rate": float(predicted_rate) if predicted_rate else None,
                "hourly_rate_was_predicted": was_predicted,
                "salary_source": salary_source,
                "hours_per_week": int(job.hours_per_week) if job.hours_per_week else 0,
                "apply_link": job.apply_link,
                "job_id": job.job_id
            })

        pair_jobs_data = []
        for match in matches.get("pairs", []):
            jobA, jobB = match.jobs[0], match.jobs[1]
            # predict missing for pair jobs as well
            def _job_salary_info(job_obj):
                o = getattr(job_obj, "hourly_rate", None)
                o_val = float(o) if o not in (None, 0, "", False) else None
                pred = None
                pred_flag = False
                if o_val is None and salary_model:
                    jfeat = {
                        "job_title": getattr(job_obj, "title", ""),
                        "required_skills": getattr(job_obj, "skills", []),
                        "experience_level": getattr(job_obj, "experience_level", ""),
                        "job_location": getattr(job_obj, "location", ""),
                        "job_type": getattr(job_obj, "job_type", ""),
                        "company_industry": getattr(job_obj, "company_industry", "")
                    }
                    pred = predict_missing_salary(jfeat, model=salary_model)
                    pred_flag = bool(pred)
                    try:
                        if pred_flag:
                            job_obj.hourly_rate_was_predicted = True
                    except Exception:
                        pass
                display = o_val if o_val is not None else (float(pred) if pred else 0)
                src = "predicted" if pred_flag else ("reported" if o_val is not None else "unknown")
                return {"original_hourly_rate": o_val, "hourly_rate": display, "predicted_hourly_rate": float(pred) if pred else None, "hourly_rate_was_predicted": pred_flag, "salary_source": src}

            a_info = _job_salary_info(jobA)
            b_info = _job_salary_info(jobB)
            pair_jobs_data.append({
                 "jobs": [
                     {
                         "title": jobA.title,
                         "company": jobA.company,
                         "location": jobA.location,
                         "original_hourly_rate": a_info["original_hourly_rate"],
                         "hourly_rate": a_info["hourly_rate"],
                         "predicted_hourly_rate": a_info["predicted_hourly_rate"],
                         "hourly_rate_was_predicted": a_info["hourly_rate_was_predicted"],
                         "salary_source": a_info["salary_source"],
                         "apply_link": jobA.apply_link,
                         "job_id": jobA.job_id
                     },
                     {
                         "title": jobB.title,
                         "company": jobB.company,
                         "location": jobB.location,
                         "original_hourly_rate": b_info["original_hourly_rate"],
                         "hourly_rate": b_info["hourly_rate"],
                         "predicted_hourly_rate": b_info["predicted_hourly_rate"],
                         "hourly_rate_was_predicted": b_info["hourly_rate_was_predicted"],
                         "salary_source": b_info["salary_source"],
                         "apply_link": jobB.apply_link,
                         "job_id": jobB.job_id
                     }
                 ],
                 "total_hours": int(match.total_hours),
                 "total_pay": float(match.total_pay),
                 "score": int(getattr(match, 'score', 0) * 100),
                 "grid": match.schedule_grid if hasattr(match, 'schedule_grid') else None
             })

        schedule_data = []
        # Basic schedule viz logic
        # Includes both Single jobs and Pair jobs
        jobs_for_schedule = []
        jobs_for_schedule.extend(matches.get("raw_top_jobs", []))
        for p in matches.get("pairs", []):
            jobs_for_schedule.extend(p.jobs)
            
        # Deduplicate by ID
        seen_sched_ids = set()
        for job in jobs_for_schedule:
            if job.job_id in seen_sched_ids: continue
            seen_sched_ids.add(job.job_id)
            
            if hasattr(job, "schedule_blocks") and job.schedule_blocks:
                 schedule_data.append({
                    "job_id": job.job_id,
                    "title": job.title,
                    "schedule": [
                        {"day": b.day, "start": format_time(b.start), "end": format_time(b.end)}
                        for b in job.schedule_blocks
                    ]
                })

        await self.websocket.send_json({
            "type": "jobs",
            "single_jobs": single_jobs_data,
            "pair_jobs": pair_jobs_data,
            "schedule_data": schedule_data
        })
        
        await self.send_message("Here are the best matches I found for you!")


# Canonical skill map for normalization and lightweight extraction fallback
SKILL_MAP = {
	"python": "Python",
	"django": "Django",
	"flask": "Flask",
	"fastapi": "FastAPI",
	"c++": "C++",
	"cpp": "C++",
	"cplusplus": "C++",
	"postgreSQL".lower(): "PostgreSQL",
	"postgresql": "PostgreSQL",
	"postgres": "PostgreSQL",
	"redis": "Redis",
	"docker": "Docker",
	"kubernetes": "Kubernetes",
	"aws": "AWS",
	"gcp": "GCP",
	"azure": "Azure",
	"react": "React",
	"vue": "Vue",
	"node": "Node.js",
	"node.js": "Node.js",
	"sql": "SQL",
	"mysql": "MySQL",
	"mongodb": "MongoDB",
	"terraform": "Terraform",
	"graphql": "GraphQL",
	"celery": "Celery",
	"pandas": "Pandas",
	"numpy": "NumPy",
	"scikit-learn": "Scikit-learn",
	"scikitlearn": "Scikit-learn",
	"rest api": "REST API",
	"ci/cd": "CI/CD",
	"ci cd": "CI/CD"
}

# Minimal EU countries list for expanded searching (can be extended)
EU_COUNTRIES = [
	"Austria","Belgium","Bulgaria","Croatia","Cyprus","Czech Republic","Denmark","Estonia",
	"Finland","France","Germany","Greece","Hungary","Ireland","Italy","Latvia","Lithuania",
	"Luxembourg","Malta","Netherlands","Poland","Portugal","Romania","Slovakia","Slovenia",
	"Spain","Sweden"
]

def _build_query_variants(profile: Dict[str, Any]) -> List[str]:
	"""
	Return prioritized query variants (role+skills, skills-only, role-only, and remote-tagged variants).
	"""
	role = (profile.get("job_role") or "").strip()
	skills = profile.get("skills", []) or []
	skills_part = " ".join(skills[:4]) if isinstance(skills, (list, tuple)) else str(skills).strip()
	out = []
	if role and skills_part:
		out.append(f"{role} {skills_part}")
	if role:
		out.append(role)
	if skills_part:
		out.append(skills_part)
	if profile.get("career_goals"):
		out.append(profile.get("career_goals"))
	# remote-tagged variants to catch postings that mention remote in title/description
	if "remote" not in " ".join(out).lower():
		for q in list(out)[:3]:
			out.append(f"{q} remote")
			out.append(f"remote {q}")
	# fallback generic roles
	if not out:
		out.append("Software Engineer")
	out.append("Software Engineer remote")
	# dedupe and return
	seen = set()
	res = []
	for q in out:
		qs = q.strip()
		if qs and qs not in seen:
			seen.add(qs)
			res.append(qs)
	return res

def _build_location_candidates(user_loc: str = "", remote_ok: bool = False) -> List[str]:
	"""
	Return prioritized location tokens including Remote, user's location, EU countries and a global fallback.
	"""
	candidates = []
	# Remote first if user is open
	if remote_ok:
		candidates.append("Remote")
	# user's normalized location
	if user_loc:
		candidates.append(user_loc)
	# if user is in EU provide EU countries as candidates
	for c in EU_COUNTRIES:
		if c not in candidates:
			candidates.append(c)
	# region tokens
	for tok in ("EU", "Europe"):
		if tok not in candidates:
			candidates.append(tok)
	# global fallback (None or empty string will be passed to discovery as no-location)
	candidates.append("")
	# dedupe preserve order
	out = []
	seen = set()
	for x in candidates:
		if x not in seen:
			seen.add(x)
			out.append(x)
	return out

# -- New endpoints for training and prediction (for testing) --
@app.post("/train_salary_model")
async def train_salary_model(payload: Dict[str, Any]):
    """
    Trigger training. Payload options:
      - {"from_db": true, "dsn": "<optional dsn>"}  -> uses DB to prepare data
      - {"csv_path": "path/to/csv"}                 -> uses CSV
      - otherwise uses ml_model/training_data.csv if present (will auto-generate if missing)
    Returns training metrics (mae/rmse) or an error.
    """
    global salary_model
    try:
        df = None
        if payload.get("from_db"):
            dsn = payload.get("dsn")
            df = prepare_training_data(dsn, out_csv="ml_model/training_data.csv")
        elif payload.get("csv_path"):
            df = pd.read_csv(payload["csv_path"])
        else:
            default_csv = "ml_model/training_data.csv"
            if not pd.io.common.file_exists(default_csv):
                # Attempt to generate synthetic training data automatically
                try:
                    from ml_model.generate_synthetic_data import generate_synthetic_dataset
                    generate_synthetic_dataset(out_csv=default_csv, n=2000, seed=42)
                    print(f"[INFO] Synthetic training data generated at {default_csv}")
                except Exception as e:
                    return JSONResponse({"status": "error", "error": f"No CSV at {default_csv} and generation failed: {e}"}, status_code=400)
            df = pd.read_csv(default_csv)

        if df is None or df.empty:
            return JSONResponse({"status": "error", "error": "No training data found."}, status_code=400)

        res = train_and_save(df)
        # reload model into memory
        salary_model = load_model()
        model_path = res.get("model_path", MODEL_PATH_DEFAULT) if isinstance(res, dict) else MODEL_PATH_DEFAULT
        model_mtime = os.path.getmtime(model_path) if os.path.exists(model_path) else None
        print(f"[OK] Salary model trained and saved to {model_path}")
        return {"status": "ok", "metrics": res, "model_path": model_path, "model_mtime": model_mtime}
    except Exception as e:
        return JSONResponse({"status": "error", "error": str(e)}, status_code=500)

@app.post("/predict_salary")
async def predict_salary(job: Dict[str, Any]):
    """
    Predict salary for a single job dict. Expected keys:
      job_title, required_skills (list or space/comma string), experience_level,
      job_location, job_type, company_industry (optional)
    """
    global salary_model
    try:
        if not salary_model:
            # attempt to load
            salary_model = load_model()
            if not salary_model:
                return JSONResponse({"status": "error", "error": "Salary model not available; train it first."}, status_code=400)

        # Normalize required_skills to list or space-separated string accepted by predictor
        skills = job.get("required_skills", "")
        if isinstance(skills, str) and "," in skills:
            job["required_skills"] = " ".join([s.strip() for s in skills.split(",") if s.strip()])

        pred = predict_missing_salary(job, model=salary_model)
        if pred is None:
            return JSONResponse({"status": "error", "error": "Prediction failed or model returned None."}, status_code=500)
        return {"status": "ok", "predicted_hourly_rate": float(pred)}
    except Exception as e:
        return JSONResponse({"status": "error", "error": str(e)}, status_code=500)

@app.websocket("/ws/chat")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()
    session = ChatSession(websocket)
    
    # Process initial empty message to trigger greeting
    try:
        await session.process_input("")
    except Exception as e:
        print("[WARN] Error during initial session greeting:", e)
        try:
            # send a minimal friendly message so client doesn't hang
            await session.send_message("Hi! Some services are currently unavailable. Please try again shortly.")
        except Exception:
            pass

    try:
        while True:
            data = await websocket.receive_text()
            await session.handle_message(data)
    except WebSocketDisconnect:
        print(f"Client disconnected")

@app.get("/")
async def get():
    return HTMLResponse(open("static/index.html", "r", encoding="utf-8").read())

@app.get("/model_status")
async def model_status():
    """Return basic model load/existence info for UI visibility."""
    exists = os.path.exists(MODEL_PATH_DEFAULT)
    mtime = os.path.getmtime(MODEL_PATH_DEFAULT) if exists else None
    loaded = bool(salary_model)
    return {"loaded": loaded, "model_path": MODEL_PATH_DEFAULT, "exists": exists, "model_mtime": mtime}

def _sanitize_static_files(root_dir: str = "static") -> Dict[str, Any]:
	"""
	Scan text files under root_dir and remove known mojibake sequences.
	Returns a summary dict with counts and modified filenames.
	"""
	seqs = ["ðŸ‘‹", "ï»¿"]  # add other sequences here if needed
	modified = []
	total_replacements = 0
	if not os.path.isdir(root_dir):
		return {"ok": False, "reason": f"{root_dir} not found", "modified": []}

	for dirpath, _, filenames in os.walk(root_dir):
		for fn in filenames:
			if not fn.lower().endswith((".html", ".htm", ".js", ".css", ".txt")):
				continue
			fp = os.path.join(dirpath, fn)
			try:
				with open(fp, "rb") as f:
					b = f.read()
				try:
					s = b.decode("utf-8")
				except Exception:
					# skip files that are not utf-8 text
					continue
				orig = s
				for seq in seqs:
					if seq in s:
						s = s.replace(seq, "")
				# remove BOM-like stray chars
				if "ï»¿" in s:
					s = s.replace("ï»¿", "")
				if s != orig:
					# backup once
					bak = fp + ".bak"
					if not os.path.exists(bak):
						with open(bak, "wb") as bf:
							bf.write(orig.encode("utf-8"))
					with open(fp, "wb") as f:
						f.write(s.encode("utf-8"))
					modified.append(fp)
					total_replacements += 1
			except Exception as e:
				print(f"[WARN] Sanitizer error on {fp}: {e}")
	return {"ok": True, "modified": modified, "replacements": total_replacements}

# Run sanitizer at startup so static files shown to users are cleaned automatically
_startup_sanitize = _sanitize_static_files("static")
if _startup_sanitize.get("modified"):
	print(f"[INFO] Sanitized {len(_startup_sanitize['modified'])} static files at startup.")
	# small listing for debug
	for m in _startup_sanitize["modified"][:10]:
		print(f"[INFO] cleaned: {m}")

@app.post("/sanitize_static")
async def sanitize_static_endpoint():
	"""Trigger static files sanitization (removes known mojibake sequences)."""
	res = _sanitize_static_files("static")
	if res.get("ok"):
		return {"status": "ok", "modified": res.get("modified", []), "replacements": res.get("replacements", 0)}
	return JSONResponse({"status": "error", "error": res.get("reason", "unknown")}, status_code=500)
